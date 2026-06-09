#!/usr/bin/env python3
"""Build the Orchestrator Air-Gap activation guide as a Word document."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "orchestrator-air-gap-activation-guide.docx"

BLUE = RGBColor(31, 77, 120)
HEADING_BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(20, 31, 43)
MUTED = RGBColor(91, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WARNING_FILL = "FFF4CE"
BORDER = "B7C4D1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_border_bottom(paragraph, color="2E74B5", size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def set_run_font(run, size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 22, DARK, 0, 6),
        ("Subtitle", 12, MUTED, 0, 14),
        ("Heading 1", 16, HEADING_BLUE, 16, 8),
        ("Heading 2", 13, HEADING_BLUE, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        if name.startswith("Heading"):
            style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("HPE Aruba EdgeConnect SD-WAN")
    set_run_font(left, size=9, color=MUTED, bold=True)
    header.add_run("\t")
    right = header.add_run("Air-Gap Activation Guide")
    set_run_font(right, size=9, color=MUTED)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    set_paragraph_border_bottom(header, color="D5DCE3", size="4")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("SD-WAN demo reference")
    set_run_font(run, size=9, color=MUTED)


def add_metadata(doc: Document) -> None:
    rows = [
        ("Audience", "Operators building or rehearsing an isolated EdgeConnect demo"),
        ("Demo target", "1 self-hosted Orchestrator and 3 EdgeConnect EC-S-P appliances"),
        ("Source", "HPE Aruba Air-Gap tab documentation and Air-Gap User Guide URL"),
        ("Generated", date.today().isoformat()),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, bold=True, color=DARK)
        value_run = p.add_run(value)
        set_run_font(value_run, color=DARK)


def add_callout(doc: Document, label: str, text: str, fill: str = WARNING_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    set_table_borders(table, color="D8C46A", size="4")
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=DARK)
    body = p.add_run(text)
    set_run_font(body, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, color=DARK)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, color=DARK)


def add_two_column_table(doc: Document, headers: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = (Inches(2.0), Inches(4.5))
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, bold=True, color=DARK)
    for left, right in rows:
        cells = table.add_row().cells
        for idx, value in enumerate((left, right)):
            cell = cells[idx]
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> None:
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Orchestrator Air-Gap Activation Guide")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Project reference for the SD-WAN demo environment")
    add_metadata(doc)
    rule = doc.add_paragraph()
    set_paragraph_border_bottom(rule)

    add_callout(
        doc,
        "Critical constraint",
        "HPE documents that once Air-Gap mode is enabled, it cannot be disabled "
        "without Silver Peak Support. Treat activation as a deliberate operator action.",
    )

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This guide summarizes the activation path for HPE Aruba Networking "
        "EdgeConnect SD-WAN Orchestrator Air-Gap deployments and records the "
        "automation guardrails for this demo project."
    )

    doc.add_heading("Official Sources", level=1)
    add_two_column_table(
        doc,
        ("Source", "Location"),
        [
            (
                "Air-Gap tab documentation",
                "https://arubanetworking.hpe.com/techdocs/sdwan/docs/orch/orchestrator/server/air-gap/",
            ),
            (
                "User guide index",
                "https://arubanetworking.hpe.com/techdocs/sdwan/user-guides/",
            ),
            (
                "Air-Gap User Guide PDF",
                "https://www.arubanetworks.com/techdocs/sdwan-PDFs/deployments/HPE-ANW-Orch-Air-gap-UG.pdf",
            ),
        ],
    )

    doc.add_heading("Demo Scope", level=1)
    add_bullets(
        doc,
        [
            "One self-hosted Orchestrator.",
            "Three EdgeConnect EC-S-P appliances.",
            "Isolated network design with manual file exchange through the Air-Gap Portal.",
            "Operator-controlled activation, licensing, appliance onboarding, and backup workflows.",
        ],
    )

    doc.add_heading("Prerequisites", level=1)
    add_bullets(
        doc,
        [
            "Air-Gap mode has been purchased and licensed.",
            "HPE Aruba Networking Operations has provisioned Air-Gap Portal access.",
            "The deployment uses self-hosted Orchestrator.",
            "The operator has Air-Gap Portal credentials and initial Orchestrator access.",
            "The team is ready to enable Air-Gap for Orchestrator and all appliances.",
            "A removable-media transfer process is approved for file movement between environments.",
        ],
    )

    doc.add_heading("Activation Flow", level=1)
    add_numbered(
        doc,
        [
            "Access the Air-Gap Portal using the HPE-provided invitation.",
            "In Orchestrator, go to Orchestrator > Orchestrator Server > Licensing > Air-Gap.",
            "Enable Air-Gap mode.",
            "Start Air-Gap Registration.",
            "Show and copy the Orchestrator registration key.",
            "Move the registration key to the Air-Gap Portal using the approved transfer process.",
            "Generate or obtain the portal response in the Air-Gap Portal.",
            "Move the portal response back to the isolated Orchestrator.",
            "Paste the portal response into Orchestrator and save it.",
            "Assign licenses in the Air-Gap Portal.",
            "Download the license file and supporting file from the Air-Gap Portal.",
            "Upload the license file and supporting file in Orchestrator with Air-Gap File Upload.",
            "Enable or register Air-Gap on each EC-S-P appliance as required by the user guide.",
            "For EC-V workflows, assign serial numbers in Configuration > Overlays & Security > Discovery > Discovered Appliances.",
        ],
    )

    doc.add_heading("Automation Guardrails", level=1)
    add_bullets(
        doc,
        [
            "Do not automate Air-Gap enablement without explicit operator confirmation.",
            "Do not store registration keys, portal responses, license files, supporting files, account keys, or serial-number assignment records in Git.",
            "Keep generated transfer manifests under an ignored runtime or output directory.",
            "Separate read-only discovery scripts from scripts that enable Air-Gap mode, upload files, assign licenses, or bind serial numbers.",
            "Prefer dry-run and checklist output unless a script is explicitly invoked with a write or apply flag.",
        ],
    )

    doc.add_heading("Project Files", level=1)
    add_two_column_table(
        doc,
        ("File", "Use"),
        [
            ("docs/orchestrator-air-gap-activation.md", "Source Markdown reference."),
            ("docs/edgeconnect-api-endpoints.md", "API endpoint source map."),
            ("docs/demo-topology.md", "Demo network topology assumptions."),
            ("inventory/demo-network.yaml", "Editable inventory for Orchestrator and EC-S-P appliances."),
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
