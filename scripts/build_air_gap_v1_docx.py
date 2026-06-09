#!/usr/bin/env python3
"""Build the v1 Air-Gap runbook DOCX from the canonical Markdown guide."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
SOURCE = REPO / "docs" / "air-gap-registration-and-ecsp-licensing-guide.md"
OUT = REPO / "docs" / "air-gap-registration-and-ecsp-licensing-guide-v1.docx"

HEADING_BLUE = RGBColor(31, 78, 121)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
BORDER = "D7DEE8"
HEADER_FILL = "E8EEF5"
NOTE_FILL = "F4F7FB"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_node = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run_node.append(r_pr)
    run_node.append(text_node)
    hyperlink.append(run_node)
    paragraph._p.append(hyperlink)


def add_runs(paragraph, text: str, *, bold_default=False) -> None:
    # Handles backtick code spans and inline URLs while keeping output simple.
    url_match = re.fullmatch(r"https?://\S+", text.strip())
    if url_match:
        add_hyperlink(paragraph, text.strip(), text.strip())
        return

    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        run.bold = bold_default
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(55, 65, 81)


def style_paragraph(paragraph, size=10.5, color=INK, before=0, after=5, line=1.08) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(size)
        run.font.color.rgb = color


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTE_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.color.rgb = HEADING_BLUE
    add_runs(p, body)
    style_paragraph(p, size=10, after=0)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in [
        ("Heading 1", 16, 12, 6),
        ("Heading 2", 12.5, 8, 4),
        ("Heading 3", 11, 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name == "Heading 1" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = HEADING_BLUE
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Air-Gap Orchestrator Registration and EC-S-P Licensing Guide | Version 1.0"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(header, size=8.5, color=MUTED, after=0)

    footer = section.footer.paragraphs[0]
    footer.text = "Keep registration keys, account keys, license files, and supporting files outside Git."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(footer, size=8.5, color=MUTED, after=0)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Air-Gap Orchestrator Registration and EC-S-P Licensing Guide")
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Version 1.0 | HPE Aruba Networking EdgeConnect SD-WAN demo runbook").bold = True
    style_paragraph(subtitle, size=11, color=MUTED, after=10)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].startswith("|"):
        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r"-+", cell.replace(" ", "")) for cell in row):
            rows.append(row)
        i += 1
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            p = cell.paragraphs[0]
            add_runs(p, value, bold_default=(r_idx == 0))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (1, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, size=8.8 if len(rows[0]) > 4 else 9.5, after=0)
    for col in table.columns:
        col.width = Inches(6.35 / len(rows[0]))


def build_docx() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_title(doc)

    i = 1  # skip Markdown title; Word title is custom.
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "Version: 1.0":
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
            style_paragraph(p, size=10.2, after=3)
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.24)
            p.paragraph_format.first_line_indent = Inches(-0.24)
            add_runs(p, line)
            style_paragraph(p, size=10.2, after=3)
        elif re.match(r"^\d+\.\d+\.", line):
            p = doc.add_paragraph()
            add_runs(p, line, bold_default=True)
            style_paragraph(p, size=10.5, color=HEADING_BLUE, before=6, after=3)
        elif line.startswith("Do not commit"):
            callout_parts = [line]
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt or nxt.startswith("#") or nxt.startswith("- ") or re.match(r"^\d", nxt) or nxt.startswith("|"):
                    break
                callout_parts.append(nxt)
                j += 1
            add_callout(doc, "Security note", " ".join(callout_parts))
            i = j
            continue
        else:
            p = doc.add_paragraph()
            add_runs(p, line)
            style_paragraph(p)
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    if not SOURCE.exists():
        sys.exit(f"Missing source Markdown: {SOURCE}")
    build_docx()
